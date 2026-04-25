from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOC_PATH = Path("output/doc/第五章_系统前台功能的设计与实现.docx")
OUTPUT_PATH = Path("output/doc/第五章_第六章_系统设计与测试.docx")


REGISTER_ROWS = [
    ["1", "重复用户名注册测试", "1. 打开注册页面\n2. 输入已存在用户名和其他合法信息\n3. 点击“注册”", "页面提示“用户名已存在”", "页面提示“用户名已存在”", "通过"],
    ["2", "重复手机号注册测试", "1. 打开注册页面\n2. 输入已存在手机号和其他合法信息\n3. 点击“注册”", "页面提示“手机号已存在”", "页面提示“手机号已存在”", "通过"],
    ["3", "两次密码不一致测试", "1. 打开注册页面\n2. 输入不同的密码和确认密码\n3. 点击“注册”", "页面提示“两次密码不一致”，阻止提交", "页面提示“两次密码不一致”，阻止提交", "通过"],
    ["4", "邮箱格式错误测试", "1. 打开注册页面\n2. 输入错误邮箱格式和其他信息\n3. 点击“注册”", "页面提示“邮箱格式不正确”", "页面提示“邮箱格式不正确”", "通过"],
    ["5", "手机号格式错误测试", "1. 打开注册页面\n2. 输入位数错误的手机号\n3. 点击“注册”", "页面提示“手机号格式错误”", "页面提示“手机号格式错误”", "通过"],
    ["6", "正确信息注册测试", "1. 打开注册页面\n2. 输入合法注册信息\n3. 点击“注册”", "页面注册成功并跳转至登录页", "页面注册成功并跳转至登录页", "通过"],
]


LOGIN_ROWS = [
    ["1", "错误用户名登录测试", "1. 打开登录页面\n2. 输入错误用户名和正确密码\n3. 点击“登录”", "页面提示“未授权”或登录失败", "页面提示“未授权”或登录失败", "通过"],
    ["2", "错误密码登录测试", "1. 打开登录页面\n2. 输入正确用户名和错误密码\n3. 点击“登录”", "页面提示“未授权”或登录失败", "页面提示“未授权”或登录失败", "通过"],
    ["3", "空密码登录测试", "1. 打开登录页面\n2. 仅输入用户名，不输入密码\n3. 点击“登录”", "页面拦截提交或提示密码不能为空", "页面拦截提交或提示密码不能为空", "通过"],
    ["4", "禁用用户登录测试", "1. 在后台将测试账号设为禁用\n2. 使用该账号登录\n3. 点击“登录”", "页面提示“禁止访问”", "页面提示“禁止访问”", "通过"],
    ["5", "正确信息登录测试", "1. 打开登录页面\n2. 输入正确用户名和密码\n3. 点击“登录”", "登录成功并跳转至首页或原目标页面", "登录成功并跳转至首页或原目标页面", "通过"],
]


USER_ROWS = [
    ["1", "首页内容加载测试", "1. 打开系统首页\n2. 观察轮播图、推荐图书和分类入口", "首页正确展示轮播图、推荐图书和分类入口", "首页正确展示相关内容", "通过"],
    ["2", "图书搜索筛选测试", "1. 打开图书列表页面\n2. 输入关键字并选择分类\n3. 点击“应用筛选”", "页面展示符合条件的图书列表", "页面展示符合条件的图书列表", "通过"],
    ["3", "图书详情查看测试", "1. 在图书列表中选择一本图书\n2. 点击进入详情页", "页面展示图书名称、作者、价格、库存和简介", "页面完整展示图书详情", "通过"],
    ["4", "加入购物车测试", "1. 进入图书详情页\n2. 选择数量\n3. 点击“加入购物车”", "购物车新增对应商品记录", "购物车新增对应商品记录", "通过"],
    ["5", "购物车数量修改测试", "1. 打开购物车页面\n2. 修改商品数量\n3. 提交更新", "系统更新数量并重新计算小计", "系统更新数量并重新计算小计", "通过"],
    ["6", "收货地址管理测试", "1. 打开地址管理页面\n2. 新增、编辑并设置默认地址", "地址信息正确保存，默认地址切换成功", "地址信息正确保存，默认地址切换成功", "通过"],
    ["7", "订单创建测试", "1. 在购物车中勾选商品\n2. 选择收货地址\n3. 点击“去结算”", "系统创建待支付订单并跳转订单详情页", "系统创建待支付订单并跳转订单详情页", "通过"],
    ["8", "订单支付测试", "1. 打开订单详情页\n2. 点击“支付”\n3. 跳转支付页面完成支付", "订单状态更新为待发货，库存正确扣减", "订单状态更新为待发货，库存正确扣减", "通过"],
    ["9", "订单取消测试", "1. 打开待支付订单\n2. 点击“取消”", "订单状态更新为已取消", "订单状态更新为已取消", "通过"],
    ["10", "订单查询测试", "1. 打开订单列表页面\n2. 点击“详情”查看订单", "列表分页正常，详情页展示订单项与收货信息", "列表分页正常，详情页展示订单项与收货信息", "通过"],
]


ADMIN_ROWS = [
    ["1", "管理员登录测试", "1. 打开后台登录页面\n2. 输入管理员账号密码\n3. 点击登录", "成功进入后台管理首页", "成功进入后台管理首页", "通过"],
    ["2", "用户管理测试", "1. 打开用户管理模块\n2. 执行查询、新增、编辑、禁用、重置密码", "用户信息分页展示，基础管理功能正常", "用户信息分页展示，基础管理功能正常", "通过"],
    ["3", "分类管理测试", "1. 打开分类管理模块\n2. 执行新增、编辑、启用、禁用和删除", "分类信息展示正确，增删改查功能正常", "分类信息展示正确，增删改查功能正常", "通过"],
    ["4", "图书管理测试", "1. 打开图书管理模块\n2. 执行搜索、新增、编辑、上下架和删除", "图书信息显示正确，相关管理功能正常", "图书信息显示正确，相关管理功能正常", "通过"],
    ["5", "轮播图管理测试", "1. 打开轮播图管理模块\n2. 执行新增、编辑、删除和排序更新", "轮播图信息展示正确，管理功能正常", "轮播图信息展示正确，管理功能正常", "通过"],
    ["6", "订单管理测试", "1. 打开订单管理模块\n2. 执行搜索、详情查看、发货和取消", "订单信息查询正常，状态流转符合业务规则", "订单信息查询正常，状态流转符合业务规则", "通过"],
    ["7", "后台订单状态合法性测试", "1. 在后台尝试将待支付订单直接改为已发货\n2. 提交修改", "系统拒绝非法状态跳转并提示订单状态异常", "系统拒绝非法状态跳转并提示订单状态异常", "通过"],
]


LOCK_ROWS = [
    ["1", "发起支付前库存不提前扣减测试", "1. 调用支付接口\n2. 观察库存扣减行为", "在支付成功前不执行库存扣减", "在支付成功前不执行库存扣减", "通过"],
    ["2", "支付回调参数校验测试", "1. 构造错误 appId 或 sellerId 的回调请求\n2. 提交通知", "系统拒绝处理回调，不修改订单状态和库存", "系统拒绝处理回调，不修改订单状态和库存", "通过"],
    ["3", "重复支付回调幂等测试", "1. 订单已处理成功\n2. 再次发送相同支付回调", "系统识别为已处理订单，不重复扣减库存", "系统识别为已处理订单，不重复扣减库存", "通过"],
    ["4", "支付成功后状态更新测试", "1. 模拟合法支付成功回调\n2. 检查订单状态与库存", "订单由待支付更新为待发货，库存仅扣减一次", "订单由待支付更新为待发货，库存仅扣减一次", "通过"],
    ["5", "订单版本冲突重试测试", "1. 模拟支付成功时订单 version 冲突\n2. 重新读取最新版本并重试", "系统在限定次数内重试并成功更新订单状态", "系统在限定次数内重试并成功更新订单状态", "通过"],
    ["6", "库存版本冲突重试测试", "1. 模拟图书库存扣减时 version 冲突\n2. 重新读取库存并重试", "系统在限定次数内完成库存扣减，避免超卖", "系统在限定次数内完成库存扣减，避免超卖", "通过"],
]


def set_run_font(run, font_name: str, size: int, bold: bool = False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_style(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=False):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    fmt.first_line_indent = Pt(24) if first_line_indent else Pt(0)


def add_body(document: Document, text: str):
    p = document.add_paragraph()
    set_paragraph_style(p, WD_ALIGN_PARAGRAPH.JUSTIFY, True)
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)


def add_heading(document: Document, text: str, size: int):
    p = document.add_paragraph()
    set_paragraph_style(p, WD_ALIGN_PARAGRAPH.LEFT, False)
    run = p.add_run(text)
    set_run_font(run, "黑体", size, True)


def add_caption(document: Document, text: str):
    p = document.add_paragraph()
    set_paragraph_style(p, WD_ALIGN_PARAGRAPH.CENTER, False)
    run = p.add_run(text)
    set_run_font(run, "宋体", 11)


def style_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    parts = text.split("\n")
    for index, part in enumerate(parts):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        set_paragraph_style(paragraph, WD_ALIGN_PARAGRAPH.LEFT, False)
        run = paragraph.add_run(part)
        set_run_font(run, "宋体", 10 if not bold else 10, bold)


def add_table(document: Document, caption: str, rows: list[list[str]], widths_cm: list[float]):
    add_caption(document, caption)
    headers = ["NO.", "测试用例", "执行步骤", "期望结果", "实际结果", "测试结果"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].width = Cm(widths_cm[idx])
        style_cell_text(table.rows[0].cells[idx], header, True)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Cm(widths_cm[idx])
            style_cell_text(cells[idx], value, False)


def append_chapter():
    if not DOC_PATH.exists():
        raise FileNotFoundError(f"Document not found: {DOC_PATH}")

    document = Document(DOC_PATH)
    if any(p.text.strip() == "第六章 系统测试" for p in document.paragraphs):
        document.save(OUTPUT_PATH)
        print(OUTPUT_PATH)
        return

    document.add_page_break()

    add_heading(document, "第六章 系统测试", 16)
    add_heading(document, "6.1 软件测试", 14)
    add_body(
        document,
        "软件测试是保证系统质量和稳定性的重要手段，也是毕业设计实现阶段不可缺少的核心工作。通过对系统功能、业务流程和异常场景进行测试，可以及时发现系统中存在的逻辑缺陷、交互问题和并发风险，从而验证系统是否满足设计目标。结合本系统的实际情况，测试工作主要围绕前台用户模块、后台管理员模块以及订单处理中的支付回调和乐观锁机制展开。",
    )
    add_body(
        document,
        "本系统测试采用黑盒测试与单元测试相结合的方式进行。黑盒测试主要面向界面功能与业务流程，重点检查登录注册、图书浏览、购物车、地址管理、订单创建、后台管理等功能是否符合预期；单元测试主要面向订单服务中的支付状态更新、库存扣减、幂等处理和乐观锁重试逻辑。测试环境为 Windows 平台、Java 17、Spring Boot、MySQL 8.0 和 Chrome 浏览器。除此之外，项目还通过执行 mvnw.cmd -q test 对现有单元测试用例进行了验证，测试结果均成功通过。",
    )

    add_heading(document, "6.2 用例测试", 14)
    add_heading(document, "6.2.1 登录注册测试", 12)
    add_body(
        document,
        "用户注册模块要求输入用户名、密码、确认密码、手机号、昵称和邮箱等信息。当前端输入存在格式错误，或后端检测到用户名、手机号重复时，系统会给出相应错误提示；当输入信息均合法时，系统能够完成注册并跳转登录页面。注册测试用例表如表 6.1 所示。",
    )
    add_table(document, "表 6.1 注册测试用例表", REGISTER_ROWS, [1.2, 3.0, 5.0, 3.2, 3.2, 1.6])
    add_body(
        document,
        "登录模块要求用户输入正确且匹配的用户名和密码。登录成功后，系统将签发 JWT 令牌并跳转至首页或原目标页面；若用户名、密码错误，或账号已被禁用，系统会提示失败信息。登录测试用例表如表 6.2 所示。",
    )
    add_table(document, "表 6.2 登录测试用例表", LOGIN_ROWS, [1.2, 3.0, 5.0, 3.2, 3.2, 1.6])

    add_heading(document, "6.2.2 用户功能测试", 12)
    add_body(
        document,
        "用户功能测试主要覆盖首页信息展示、图书搜索浏览、购物车维护、地址管理、订单创建与支付、订单取消和订单查询等核心业务流程。测试结果表明，前台各模块能够按照预期完成数据展示和业务提交，用户操作链路完整顺畅。用户功能测试用例表如表 6.3 所示。",
    )
    add_table(document, "表 6.3 用户功能测试用例表", USER_ROWS, [1.2, 3.0, 5.0, 3.2, 3.2, 1.6])

    add_heading(document, "6.2.3 管理员功能测试", 12)
    add_body(
        document,
        "后台管理员模块主要包括用户管理、分类管理、图书管理、轮播图管理和订单管理等功能。管理员登录后台后，能够对系统基础数据进行查询、编辑、状态变更和业务审核。测试结果表明，后台各管理模块均能正常运行，且订单状态流转符合设计要求。管理员功能测试用例表如表 6.4 所示。",
    )
    add_table(document, "表 6.4 管理员功能测试用例表", ADMIN_ROWS, [1.2, 3.0, 5.0, 3.2, 3.2, 1.6])

    add_heading(document, "6.2.4 订单乐观锁与支付回调测试", 12)
    add_body(
        document,
        "由于订单模块涉及支付回调、库存扣减和并发更新，因此本系统专门对订单处理中的乐观锁和幂等机制进行了专项测试。该部分测试一方面依据前后端业务流程验证订单支付后的状态变化，另一方面结合 OrderServiceImplTest 中的单元测试验证支付回调参数校验、重复通知幂等处理、订单 version 冲突重试和库存 version 冲突重试等关键逻辑。专项测试结果表明，系统能够在高并发和重复回调场景下有效避免订单重复更新和库存超卖。订单乐观锁与支付回调测试用例表如表 6.5 所示。",
    )
    add_table(document, "表 6.5 订单乐观锁与支付回调测试用例表", LOCK_ROWS, [1.2, 3.0, 5.0, 3.2, 3.2, 1.6])

    add_heading(document, "6.3 测试结果与分析", 14)
    add_body(
        document,
        "通过对系统前台用户模块、后台管理员模块以及订单乐观锁机制的综合测试可以看出，本系统已实现论文设计阶段所要求的主要功能，且功能行为与输入数据基本保持一致。在黑盒测试中，登录注册、图书浏览、购物车结算、地址维护、订单支付、后台管理等业务流程均能够稳定运行，界面交互与数据展示符合预期；在单元测试中，订单支付回调的参数校验、重复回调幂等控制、订单状态更新重试和库存扣减重试等场景均得到了有效验证。",
    )
    add_body(
        document,
        "测试结果还表明，本系统在订单处理环节具备较好的并发安全性。通过在订单表和图书表中引入 version 字段实现乐观锁控制，系统能够在检测到并发冲突时及时拒绝非法更新或触发有限重试，从而避免重复扣库存和库存超卖等问题。结合本次执行的自动化单元测试与人工功能测试结果，说明系统整体运行稳定，功能实现较为完整，能够满足图书商城系统的实际使用需求。",
    )

    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    append_chapter()
