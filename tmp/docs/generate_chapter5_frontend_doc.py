from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_PATH = Path("output/doc/第五章_系统前台功能的设计与实现.docx")


SECTIONS = [
    (
        "heading1",
        "5.1 系统前台总体设计与实现",
    ),
    (
        "body",
        "本系统采用前后端分离的实现方式，前台模块面向普通用户，主要完成图书浏览、用户登录注册、购物车维护、收货地址管理、订单创建、支付和订单查询等业务操作。前端基于 Vue 3、Vite、Pinia 和 Axios 实现，后端基于 Spring Boot、MyBatis 与 MySQL 提供 RESTful 接口，前后端通过 JSON 数据进行异步交互。该实现方式既保证了界面的交互性，也便于后续对业务逻辑和页面展示进行独立维护。",
    ),
    (
        "body",
        "系统前台主页面由导航栏、首页轮播图、分类快捷入口、推荐图书和畅销榜等区域组成。未登录状态下，用户可以直接浏览首页、图书列表和图书详情等公共内容；登录成功后，系统在导航栏中提供用户中心、购物车、订单管理、地址管理和退出登录等功能入口，并通过路由守卫限制未登录用户访问需要认证的页面，从而保证前台业务流程的完整性与安全性。系统前台主页面如图 5.1 所示。",
    ),
    (
        "figure",
        "图 5.1 系统前台主页面（此处插入首页截图）",
    ),
    (
        "heading2",
        "5.1.1 前台用户登录注册功能的设计与实现",
    ),
    (
        "body",
        "用户登录注册功能是前台业务的起点。只有用户完成注册并成功登录后，才能继续使用购物车、收货地址、订单支付和订单查询等个性化服务。在本系统中，注册页面由 Register.vue 实现，登录页面由 Login.vue 实现，前端通过 Axios 封装的 request 模块向后端发送异步请求，后端通过 FrontAuthController 接收请求，再由 UserServiceImpl 完成用户校验、密码加密和令牌生成等处理。",
    ),
    (
        "body",
        "在用户注册功能中，用户需要在前端页面填写用户名、密码、确认密码、手机号、昵称和邮箱等信息。为了提升输入数据的准确性，前端在用户失焦时即进行格式校验，例如用户名要求为 4 至 20 位字母、数字或下划线，密码要求同时包含字母和数字，手机号要求为 11 位数字，邮箱字段则按标准邮箱格式进行校验。当前端校验通过后，页面调用 frontAuthApi.register 方法，以 AJAX 的方式向 /api/front/auth/register 发送 POST 请求。后端接收到请求后，首先依据 RegisterRequest 中的注解规则进行参数合法性校验，然后由 UserServiceImpl.register 方法判断用户名和手机号是否重复；若校验通过，则调用 PasswordEncoder 对密码进行加密，再通过 UserMapper 将用户信息写入数据库，最终返回包含 token、userId、username 和 role 的认证结果。用户注册流程如图 5.2 所示。",
    ),
    (
        "figure",
        "图 5.2 用户注册流程图（可根据前后端调用关系补充时序图）",
    ),
    (
        "body",
        "在用户登录功能中，用户输入用户名和密码后提交表单，前端调用 frontAuthApi.login 接口向 /api/front/auth/login 发送请求。后端由 FrontAuthController.login 方法接收登录参数，再调用 UserServiceImpl.login 完成身份认证。该方法首先通过 UserMapper 查询用户名对应的用户记录，然后使用 PasswordEncoder.matches 对密码进行比对，同时判断用户是否处于禁用状态。当用户名和密码校验通过后，系统利用 JwtUtil 生成登录令牌，并将用户编号与角色信息写入 Token 的 claims 中。前端接收到返回结果后，会将 token、角色、用户编号和用户名分别写入本地存储和 Pinia 状态管理模块，并在后续请求中通过 Axios 请求拦截器自动携带 Authorization 头信息。用户登录成功后，系统会根据 redirect 参数跳转到原目标页面，若认证失败则通过统一消息提示显示错误原因。用户登录界面和注册界面可分别如图 5.3、图 5.4 所示。",
    ),
    (
        "figure",
        "图 5.3 用户注册界面（此处插入注册页面截图）",
    ),
    (
        "figure",
        "图 5.4 用户登录界面（此处插入登录页面截图）",
    ),
    (
        "heading2",
        "5.1.2 图书浏览与检索功能的设计与实现",
    ),
    (
        "body",
        "图书浏览与检索功能是前台最核心的展示功能之一，主要包括首页推荐图书展示、图书列表分页查询、按分类筛选、按关键字模糊搜索和图书详情查看等功能。首页 Home.vue 页面在组件加载时会并发调用轮播图接口、图书列表接口和分类列表接口，将获取的数据分别渲染到轮播区、分类快捷入口和推荐图书区域中，从而保证首页内容的完整展示效果。",
    ),
    (
        "body",
        "图书列表页面由 BookList.vue 实现。用户可以在页面中输入关键字，或按分类、价格区间等条件进行筛选。前端通过 frontBookApi.list 调用 /api/front/books 接口，将 page、pageSize、categoryId、keyword、minPrice 和 maxPrice 等查询参数提交到后端。后端的 FrontBookController.list 方法负责接收这些参数，再由 BookService 和 BookMapper 组合查询条件，仅返回处于上架状态且未被软删除的图书数据，同时对结果进行分页处理。查询完成后，系统将 total 和 list 返回前端，由页面完成卡片式列表渲染与分页栏展示。图书列表界面如图 5.5 所示。",
    ),
    (
        "figure",
        "图 5.5 图书列表与筛选界面（此处插入图书列表截图）",
    ),
    (
        "body",
        "当用户点击某本图书时，系统会进入图书详情页面 BookDetail.vue，并通过 frontBookApi.detail 向 /api/front/books/{id} 发起请求。后端根据图书编号查询书名、作者、出版社、ISBN、价格、库存和简介等信息，再返回给前端展示。详情页中还提供数量选择器与“加入购物车”按钮，用户可在此基础上进一步发起购买流程。通过将列表查询和详情查询相结合，系统实现了从信息展示到业务操作的自然衔接。",
    ),
    (
        "heading2",
        "5.1.3 购物车与订单创建功能的设计与实现",
    ),
    (
        "body",
        "购物车模块负责承接用户的购买意图，并为订单创建提供数据来源。用户在图书详情页面选择购买数量后，点击“加入购物车”按钮，前端会调用 frontCartApi.add 接口，将 bookId 和 quantity 发送到 /api/front/cart。后端由 FrontCartController.add 方法接收请求，根据当前登录用户编号调用 CartService.add 完成购物车记录新增或数量更新，同时校验图书是否存在、是否上架以及库存是否充足。该模块使用户能够先集中管理拟购买商品，再统一进行结算操作。",
    ),
    (
        "body",
        "购物车页面由 Cart.vue 实现，页面以表格形式展示商品名称、单价、数量、小计以及删除按钮，并通过复选框支持用户选择多个商品同时结算。当前端修改商品数量时，会调用 /api/front/cart/{id} 接口同步更新数据库；删除商品时，则调用 DELETE 接口移除对应记录。页面底部实时计算已选商品总金额，并提供“去结算”按钮触发下单流程。",
    ),
    (
        "body",
        "在订单创建功能中，用户点击“去结算”后，前端首先读取地址列表并弹出地址选择框；用户确认地址后，系统将所选地址编号和购物车编号列表封装为 OrderCreateRequest，通过 frontOrderApi.create 接口提交到 /api/front/orders。后端由 FrontOrderController.create 接收请求，并在 OrderServiceImpl.create 方法中完成订单创建逻辑。该方法首先校验地址是否属于当前用户，再查询购物车记录并逐一校验图书是否存在、是否上架以及库存是否足够，然后计算订单总金额、生成订单编号、写入订单主表和订单明细表，并删除对应购物车记录。需要强调的是，本系统在创建订单阶段并不立即扣减库存，而是仅完成订单数据落库和状态初始化，将订单状态置为“待支付”，这样可以避免大量未支付订单长期占用库存。购物车与订单创建界面如图 5.6 所示。",
    ),
    (
        "figure",
        "图 5.6 购物车结算与订单创建界面（此处插入购物车截图）",
    ),
    (
        "heading2",
        "5.1.4 收货地址管理功能的设计与实现",
    ),
    (
        "body",
        "收货地址管理功能用于维护订单配送所需的联系人和地址信息。地址页面由 AddressManage.vue 实现，支持地址列表查询、新增、编辑、删除和设置默认地址等操作。前端通过统一的弹窗表单收集收件人、联系电话、省、市、区和详细地址，并通过 frontAddressApi 与后端进行交互。地址数据在订单创建时会被进一步使用，因此该模块在前台业务链路中起到了基础支撑作用。",
    ),
    (
        "body",
        "在实现过程中，当前端页面加载时会调用 /api/front/user/addresses 接口查询当前用户的全部地址信息，并在表格中展示。若用户选择新增地址，页面会弹出空白表单并调用 POST 接口写入数据库；若用户选择编辑地址，则先回填已有地址信息，再调用 PUT 接口完成更新；若用户点击“设默认”，系统会调用 /api/front/user/addresses/{id}/default 接口，后端将该地址设为默认地址并自动取消同一用户其他地址的默认状态。该设计既满足了用户对多地址管理的需求，也保证了订单结算时默认地址的可快速选取。地址管理界面如图 5.7 所示。",
    ),
    (
        "figure",
        "图 5.7 收货地址管理界面（此处插入地址管理截图）",
    ),
    (
        "heading2",
        "5.1.5 订单支付、订单查询与取消功能的设计与实现",
    ),
    (
        "body",
        "订单支付功能建立在订单创建成功的基础之上。前端订单列表页面 OrderList.vue 和订单详情页面 OrderDetail.vue 都提供“支付”和“取消”两个操作按钮，其中只有状态为待支付的订单才允许继续支付或取消。用户点击支付按钮后，前端调用 frontOrderApi.pay 接口，向 /api/front/orders/{id}/pay 发送 POST 请求。与普通 JSON 接口不同，该接口返回的是支付宝支付表单的 HTML 内容，前端将其插入页面后自动提交表单，从而跳转到支付宝支付页面完成付款。",
    ),
    (
        "body",
        "后端在 OrderServiceImpl.pay 方法中并不直接修改订单状态，而是先验证订单归属、当前状态以及订单项对应图书的有效性，然后构造支付宝支付请求并返回表单内容。真正的支付成功处理由 AlipayController 中的 notify 和 returnUrl 两个接口完成：notify 用于接收异步通知，returnUrl 用于同步回跳页面。系统在校验支付宝签名、应用编号、卖家编号和支付金额都正确后，才会继续执行支付成功后的业务处理。该设计使支付接口与订单变更逻辑相分离，降低了前端直接修改订单状态的风险。",
    ),
    (
        "body",
        "在订单查询功能中，前端通过 /api/front/orders 接口分页获取订单列表，并在页面中展示订单号、金额、状态和创建时间等信息；点击详情按钮后，再通过 /api/front/orders/{id} 获取当前订单的完整信息，包括订单条目、收货人、联系电话和收货地址等内容。对于待支付订单，用户还可以通过 /api/front/orders/{id}/cancel 接口主动取消订单，后端在确认当前状态为待支付后将订单状态更新为已取消。订单列表与订单详情界面如图 5.8、图 5.9 所示。",
    ),
    (
        "figure",
        "图 5.8 前台订单列表界面（此处插入订单列表截图）",
    ),
    (
        "figure",
        "图 5.9 前台订单详情界面（此处插入订单详情截图）",
    ),
    (
        "heading2",
        "5.1.6 订单处理中的乐观锁设计与实现",
    ),
    (
        "body",
        "订单处理是本系统前台功能实现中的重点和难点。为了在支付回调、库存扣减和订单状态流转过程中保证数据一致性，系统在 order 表和 book 表中都引入了 version 字段，并采用基于版本号的乐观锁机制控制并发更新。新增记录时 version 默认值为 0，每次更新成功后 version 自动加 1；更新 SQL 语句必须同时满足主键匹配和版本号匹配两个条件，若影响行数为 0，则说明当前数据已经被其他事务修改，系统需要执行冲突处理逻辑。",
    ),
    (
        "body",
        "在支付成功后的订单状态更新过程中，OrderServiceImpl 会调用 markPaymentSuccess 方法完成从“待支付”到“待发货”的状态变更。该方法执行的核心 SQL 同时校验订单编号、原始状态和版本号，只有当订单仍处于待支付状态且版本号未变化时，才允许写入 pay_time、trade_no 并将 version 加 1。若更新失败，系统会重新读取最新订单状态：如果订单已经被成功处理，则直接返回成功结果；如果订单仍为待支付但版本号已变化，则系统会在限定次数内重新尝试更新。通过这种方式，可以避免用户重复支付或支付平台重复通知时造成订单状态被重复修改的问题。",
    ),
    (
        "body",
        "在库存扣减过程中，系统采用与订单状态更新相同的版本号控制策略。支付成功后，OrderServiceImpl 会遍历订单中的每个商品，调用 decreaseStock 方法读取图书当前库存和版本号，再执行带有“version = ?”和“stock >= ?”条件的更新语句。只有当当前库存足够且版本号一致时，数据库才会真正扣减库存并更新版本号；如果并发请求导致版本不匹配，系统会再次读取最新图书数据并重试。在多用户同时购买同一本图书的场景下，该设计能够显著降低超卖风险，并避免使用悲观锁带来的长时间数据库锁等待问题。",
    ),
    (
        "body",
        "除了版本号控制外，系统还对支付回调场景进行了幂等处理。当订单已经处于“待发货”状态且已记录支付时间时，系统会认为该支付已被成功处理，后续重复通知将直接返回成功，不再重复扣减库存。与此同时，系统将乐观锁冲突统一封装为业务异常，并返回“数据已变更，请刷新后重试”的提示信息。对于支付链路中的关键步骤，系统还设置了最多 3 次的有限重试，以兼顾高并发场景下的成功率与系统可控性。由此可见，乐观锁机制不仅提高了订单处理的安全性，也增强了系统在并发场景下的稳定性与工程可用性。",
    ),
    (
        "body",
        "综合来看，本系统前台订单模块采用了“创建订单不占库存、支付成功后再扣库存”的实现策略，并在支付成功处理与库存扣减环节引入乐观锁机制，从而较好地解决了重复回调、并发更新和库存超卖等问题。这一设计既符合网上商城业务的通行实现方式，也体现了系统在详细设计与实现阶段对事务一致性和并发安全的充分考虑。",
    ),
]


def set_run_font(run, font_name: str, size: int, bold: bool = False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_line_spacing(paragraph):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)


def add_body_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24)
    set_paragraph_line_spacing(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 12, False)


def add_heading(document: Document, text: str, level: int):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_line_spacing(paragraph)
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, "黑体", 14, True)
    else:
        set_run_font(run, "黑体", 12, True)


def add_figure_caption(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_line_spacing(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 11, False)


def build_document():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run("第五章 系统前台功能的设计与实现")
    set_run_font(title_run, "黑体", 16, True)

    for kind, text in SECTIONS:
        if kind == "heading1":
            add_heading(document, text, 1)
        elif kind == "heading2":
            add_heading(document, text, 2)
        elif kind == "figure":
            add_figure_caption(document, text)
        else:
            add_body_paragraph(document, text)

    document.add_section(WD_SECTION.CONTINUOUS)
    return document


def main():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
